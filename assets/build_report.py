"""
Build the Atlanta Robotics Word document project report.
Run: python3 assets/build_report.py
Output: assets/Atlanta_Robotics_DataCenter_DigitalTwin_Report.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__),
                   "Atlanta_Robotics_DataCenter_DigitalTwin_Report.docx")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Style helpers ──────────────────────────────────────────────────────────────
NVIDIA_GREEN = RGBColor(0x76, 0xB9, 0x00)
DARK_BLUE    = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT_BLUE  = RGBColor(0x00, 0x70, 0xC0)
GREY         = RGBColor(0x59, 0x59, 0x59)

def set_run_color(run, color):
    run.font.color.rgb = color

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = NVIDIA_GREEN
    p.runs[0].font.size = Pt(20)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = ACCENT_BLUE
    p.runs[0].font.size = Pt(15)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = DARK_BLUE
    p.runs[0].font.size = Pt(13)
    return p

def body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    return p

def numbered(text, level=0):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F0F0F0')
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'E8F5E9')
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run("📝  " + text)
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '76B900')
        tcPr.append(shd)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx+1].cells
        fill = 'F9F9F9' if r_idx % 2 == 0 else 'FFFFFF'
        for i, cell_text in enumerate(row_data):
            cells[i].text = str(cell_text)
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            tc = cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table

def page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Data Center Digital Twin")
run.font.size = Pt(32)
run.bold = True
run.font.color.rgb = NVIDIA_GREEN

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AI World Model on Google Cloud")
run.font.size = Pt(22)
run.font.color.rgb = ACCENT_BLUE

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("─" * 50)
run.font.color.rgb = NVIDIA_GREEN

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "A Complete Step-by-Step Tutorial for College Students\n"
    "Atlanta Robotics — Tutorial Series\n\n"
    "Covering: Digital Twin Technology · Synthetic Data Generation\n"
    "Temporal Transformer Architecture · Google Cloud Deployment\n"
    "Vertex AI Training · Live Failure Prediction"
)
run.font.size = Pt(12)
run.font.color.rgb = GREY

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2026")
run.font.size = Pt(14)
run.bold = True

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════
h1("Table of Contents")
toc_entries = [
    ("1", "Executive Summary", "3"),
    ("2", "Problem Statement", "3"),
    ("3", "What Is a Digital Twin?", "4"),
    ("4", "Understanding USD — Universal Scene Description", "5"),
    ("5", "System Architecture", "6"),
    ("6", "Phase-by-Phase Implementation Guide", "7"),
    ("  6.1", "Phase 1 — Install gcloud CLI", "7"),
    ("  6.2", "Phase 2 — GCP Infrastructure Setup", "7"),
    ("  6.3", "Phase 3 — Upload USD Assets to GCS", "8"),
    ("  6.4", "Phase 4 — Build and Push Docker Container", "9"),
    ("  6.5", "Phase 5 — Deploy to GPU VM and Launch Streaming", "9"),
    ("  6.6", "Phase 6 — Synthetic Failure Data Generation", "10"),
    ("  6.7", "Phase 7 — Temporal Transformer World Model", "11"),
    ("  6.8", "Phase 8 — Vertex AI Training Job", "14"),
    ("  6.9", "Phase 9 — Inference Integration", "15"),
    ("7", "Getting the USD Assets (Students)", "15"),
    ("8", "Training Results and Model Performance", "16"),
    ("9", "Live Inference Results", "17"),
    ("10", "Exercises for Students", "18"),
    ("11", "Key Concepts Glossary", "19"),
    ("12", "Cloud Cost Summary", "20"),
    ("13", "Common Errors and Fixes", "21"),
]
for num, title, pg in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    dots = "." * max(1, 65 - len(num) - len(title))
    run = p.add_run(f"  {num}  {title} {dots} {pg}")
    run.font.size = Pt(10)
    if not num.startswith(" "):
        run.bold = True

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Executive Summary")
body(
    "This tutorial guides college students at Atlanta Robotics through building a complete, "
    "production-grade AI system from scratch. Starting with a 9.6 GB NVIDIA data center "
    "digital twin (a photorealistic 3D model of a real server room), students simulate hardware "
    "failure scenarios, generate a labeled sensor dataset, train a Temporal Transformer neural "
    "network to predict failures before they happen, and deploy the trained model as a live "
    "inference endpoint on Google Cloud."
)
body(
    "By completing this tutorial you will have hands-on experience with: NVIDIA Omniverse and "
    "USD 3D scene composition, Google Cloud Platform (GCS, Compute Engine, Artifact Registry, "
    "Vertex AI), PyTorch Transformer architecture for timeseries, synthetic data generation, "
    "Docker containerization, and cloud ML deployment. These are exactly the skills used by "
    "robotics and AI teams at companies like NVIDIA, Google, and Amazon."
)
note("Estimated total cloud cost per student: ~$11. Estimated hands-on time: 4–6 hours.")

# ══════════════════════════════════════════════════════════════════════════════
# 2. PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Problem Statement")
body(
    "Modern data centers contain thousands of servers, storage arrays, switches, and cooling "
    "units. A single unexpected hardware failure can cause downtime costing $100,000 or more "
    "per hour. The goal of predictive maintenance AI is to detect imminent failures hours or "
    "days before they occur, allowing engineers to replace components during scheduled "
    "maintenance windows instead of emergency outages."
)
h2("2.1  The Data Problem")
body(
    "Training a machine learning model to predict hardware failures requires examples of both "
    "normal operation and failures. In a real data center with 5,000 servers, you might see "
    "10–50 failures per year — a failure rate of roughly 0.01%. Training on data this imbalanced "
    "produces a model that simply predicts 'normal' for everything and achieves 99.99% accuracy, "
    "while being completely useless for alerting."
)
body("Additional complications:")
bullet("Real failures are dangerous to allow intentionally — you cannot stress-test live hardware.")
bullet("Historical failure logs are confidential and hard to obtain from data center operators.")
bullet("Each failure type has a different sensor signature — you need labeled examples of each.")

h2("2.2  The Solution: Synthetic Data from a Digital Twin")
body(
    "The solution is to build a virtual replica of the data center — a digital twin — and "
    "simulate failure scenarios inside it. We can generate thousands of labeled failure events "
    "at no cost and no risk to real hardware. The model trains on this synthetic data and "
    "deploys against real sensor streams. This technique is used in production by Siemens, "
    "General Electric, NVIDIA, and many other industrial companies."
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. WHAT IS A DIGITAL TWIN
# ══════════════════════════════════════════════════════════════════════════════
h1("3. What Is a Digital Twin?")
body(
    "A digital twin is a virtual replica of a physical system that: (1) looks like the real "
    "thing — accurate 3D geometry, materials, and layout; (2) behaves like the real thing — "
    "physics simulation and sensor models; and (3) can be synchronized with real sensor data "
    "from the physical asset."
)
body(
    "In this project, the physical system is the NVIDIA DataHall — a real data center floor "
    "housing hundreds of DGX A100 servers, QM8700 InfiniBand switches, 42U server racks, "
    "power distribution units, and precision air cooling. The virtual replica is "
    "DataHall_Full_01.usd, a 9.6 GB file containing every rack, cable tray, cooling unit, "
    "and floor tile in the room."
)

h2("3.1  The Digital Twin Pipeline")
body("The complete pipeline from digital twin to AI prediction:")
numbered("Physical data center — 5-minute sensor readings from real DCIM systems.")
numbered("Digital twin (USD stage) — simulates failures using 3D physics and sensor models.")
numbered("Synthetic sensor dataset (CSV) — 414,720 labeled rows exported by the simulation.")
numbered("Failure prediction model (PyTorch) — trained on the synthetic dataset.")
numbered("Real-time alerts in the digital twin — failing racks glow red in the 3D viewport.")

note(
    "The key insight: the model is trained on simulated data but deployed against real data. "
    "This works because the simulation accurately models how sensors respond during failures."
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. USD FORMAT
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Understanding USD — Universal Scene Description")
body(
    "USD (Universal Scene Description) was created by Pixar Animation Studios for film "
    "production and adopted by NVIDIA for industrial simulation. It is now the standard "
    "3D interchange format for the Omniverse platform."
)

h2("4.1  Composition — Why the Root File Is Only 31 KB")
body(
    "The root file DataHall_Full_01.usd is only 31 KB, but the entire asset tree is 9.6 GB "
    "across 1,822 files. This is because USD uses composition: the root file contains only "
    "references (pointers) to sub-assets. When Omniverse loads the stage, it follows all "
    "references and assembles the full scene in GPU memory."
)
body("Think of it exactly like an HTML page: the HTML file is tiny but references CSS files, "
     "images, and JavaScript that together may be gigabytes in size.")

code_block(
    "DataHall_Full_01.usd  (31 KB — the root/index file)\n"
    "  references →  Racks/Rack_42U_A/Rack_42U_A_01.usd\n"
    "                    references →  SubUSDs/Rack_42U_inst.usdc\n"
    "                    references →  SubUSDs/DataHall.material.usd\n"
    "  references →  DGX_Servers/A100/DGX_A100_01.usd\n"
    "                    references →  SubUSDs/dgx_a100_base.usd\n"
    "                    references →  textures/dgx_chassis_color.png  (300 MB)\n"
    "  references →  Network_Switches/QM8700/QM8700_01.usd\n"
    "  ...and so on for every rack, switch, PDU, and facility element"
)

h2("4.2  Instancing — Why 48 Servers Don't Cost 48× the Memory")
body(
    "USD supports instancing: one DGX A100 model (~800 MB) can be referenced 48 times "
    "rather than copied 48 times. Omniverse loads the geometry into GPU memory once, then "
    "applies 48 different transformation matrices to position it in each rack slot. This "
    "is how the DataHall renders 48 identical servers without 38.4 GB of duplicate geometry."
)

h2("4.3  USD File Formats")
add_table(
    ["Extension", "Format", "Description"],
    [
        [".usd",  "Either",        "Container format — may be text or binary"],
        [".usda", "Text (ASCII)",  "Human-readable, good for debugging and editing"],
        [".usdc", "Binary (Crate)","Compressed binary, fast to load, smaller file size"],
        [".usdz", "Zip archive",   "Self-contained package — all sub-assets bundled inside"],
    ],
    col_widths=[1.2, 1.8, 4.5]
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
h1("5. System Architecture")
body(
    "The system has three main environments: the instructor's local machine (which builds "
    "and uploads assets), Google Cloud Platform (which hosts the GPU VM, container registry, "
    "storage, and Vertex AI), and the student's local machine (which downloads assets and "
    "trains the model)."
)

code_block(
    "LOCAL MACHINE (instructor)\n"
    "  kit-app-template  ──► repo.sh build ──► Docker image\n"
    "  DataHall_Full_01.usd ───────────────► GCS bucket\n"
    "         │                                    │\n"
    "         │  docker push          gsutil cp    │\n"
    "         ▼                                    ▼\n"
    "GOOGLE CLOUD PLATFORM\n"
    "  Artifact Registry ◄── Docker image\n"
    "  GCS Bucket\n"
    "    ├── Datacenter_NVD@10012/  (9.6 GB USD assets)\n"
    "    ├── training-data/sensor_timeseries.csv\n"
    "    └── models/best_model.pt\n"
    "  GPU VM  (g2-standard-8 + NVIDIA L4)\n"
    "    └── Kit App Container ──► Port 8011 (browser WebRTC stream)\n"
    "  Vertex AI\n"
    "    ├── Training Job  (A100 GPU)\n"
    "    └── Endpoint  (live failure probability API)\n"
    "         │\n"
    "         │  JSON: {'1h': 0.92, '6h': 0.87, '24h': 0.71}\n"
    "         ▼\n"
    "BROWSER  http://VM-IP:8011\n"
    "  Failing racks highlighted red in 3D viewport"
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. PHASE-BY-PHASE GUIDE
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Phase-by-Phase Implementation Guide")

# Phase 1
h2("6.1  Phase 1 — Install gcloud CLI")
body(
    "gcloud is Google's command-line tool that controls all GCP services. It must be "
    "installed before any other phase. The script handles the full APT repository setup."
)
code_block("bash deploy/01_install_gcloud.sh\n\n"
           "# After install, manually authenticate:\n"
           "gcloud auth login\n"
           "gcloud config set project YOUR_PROJECT_ID\n"
           "gcloud auth application-default login  # for Python SDKs")
bullet("Checkpoint: gcloud --version prints a version number.")

# Phase 2
h2("6.2  Phase 2 — GCP Infrastructure Setup")
body("This phase creates all cloud resources needed for the rest of the tutorial.")
code_block("source deploy/config.env\nbash deploy/02_gcp_setup.sh")
body("The script performs the following steps:")
numbered("Enables four GCP APIs: Compute Engine, Artifact Registry, Vertex AI, Cloud Storage.")
numbered("Creates GCS bucket gs://PROJECT-omniverse-assets/ for USD assets and training data.")
numbered("Creates Artifact Registry repository omniverse-kit for Docker images.")
numbered(
    "Creates a g2-standard-8 VM (8 vCPU, 32 GB RAM, 1× NVIDIA L4) with Ubuntu 22.04. "
    "A startup script runs on first boot to install NVIDIA drivers (550 series), the "
    "NVIDIA Container Toolkit (so Docker can use the GPU), and GCS Fuse (to mount the "
    "asset bucket as a local folder)."
)
numbered(
    "Opens firewall ports: TCP 8011 (Kit HTTP), TCP 8012 (Kit WebSocket), "
    "UDP 49100–49200 (WebRTC media stream)."
)

add_table(
    ["Resource", "Type", "Cost"],
    [
        ["GCS Bucket",        "Object storage",    "~$0.02/GB/month"],
        ["Artifact Registry", "Docker registry",   "~$0.10/GB/month"],
        ["GPU VM (L4)",       "Compute Engine",    "~$0.40/hour"],
        ["Vertex AI endpoint","Prediction serving", "~$1.50/hour (T4)"],
    ],
    col_widths=[2.5, 2.5, 2.5]
)

# Phase 3
h2("6.3  Phase 3 — Upload USD Assets to GCS")
body(
    "The 9.6 GB DataHall USD stage is uploaded from the instructor's local disk to GCS "
    "using parallel gsutil transfers. This typically takes 10–30 minutes depending on "
    "internet connection speed."
)
code_block("source deploy/config.env\nbash deploy/03_upload_assets.sh")
body(
    "The script uses gsutil -m cp -r which uploads multiple files simultaneously using "
    "parallel threads. The -z usd,usda,usdc flag enables gzip compression for text-based "
    "USD files, reducing transfer size by ~40%."
)
note(
    "Students do NOT run this script. The instructor runs it once, then grants students "
    "read-only access using deploy/instructor_grant_access.sh."
)

# Phase 4
h2("6.4  Phase 4 — Build and Push Docker Container")
body(
    "The NVIDIA Kit application (configured USD Explorer pointing at the DataHall stage) "
    "is compiled and packaged into a Docker container image, then pushed to Artifact Registry."
)
code_block(
    "source deploy/config.env\nbash deploy/04_build_and_push.sh\n\n"
    "# Internally runs:\n"
    "./repo.sh build              # compiles Kit extensions\n"
    "./repo.sh package_container  # wraps into Docker image\n"
    "docker push IMAGE_URI        # uploads to Artifact Registry"
)
body(
    "The resulting image is ~15–25 GB (Kit includes CUDA libraries, RTX renderer, and all "
    "extension bundles). Storing it in Artifact Registry in the same region as the VM "
    "eliminates egress costs and reduces pull time from 30+ minutes to ~3 minutes."
)

# Phase 5
h2("6.5  Phase 5 — Deploy to GPU VM and Launch Streaming")
body(
    "This script SSHes into the GPU VM, mounts the GCS bucket via GCS Fuse, pulls the "
    "container, and starts Kit with GPU passthrough and WebRTC streaming enabled."
)
code_block(
    "source deploy/config.env\nbash deploy/05_deploy_vm.sh\n\n"
    "# The script runs this on the VM:\n"
    "gcsfuse --implicit-dirs PROJECT-omniverse-assets /mnt/assets\n"
    "docker run -d --gpus all \\\n"
    "    -p 8011:8011 -p 8012:8012 -p 49100-49200:49100-49200/udp \\\n"
    "    -v /mnt/assets:/mnt/assets:ro \\\n"
    "    IMAGE_URI \\\n"
    "    --/app/auto_load_usd=/mnt/assets/.../DataHall_Full_01.usd \\\n"
    "    --/app/streaming/webrtc/enabled=true"
)
body(
    "After the container starts (~60 seconds for Kit to initialize), open "
    "http://VM-IP:8011 in Chrome or Firefox. The 3D DataHall scene streams live via "
    "WebRTC — you can orbit, zoom, and navigate the data center in real time."
)

page_break()

# Phase 6
h2("6.6  Phase 6 — Synthetic Failure Data Generation")
body(
    "This is the heart of the digital twin approach: using the 3D model as the basis "
    "for generating synthetic training data. The script simulates four failure types "
    "across 48 racks over 30 days at 5-minute sensor sampling intervals."
)
code_block(
    "# Standalone mode (system Python — works without Isaac Sim):\n"
    "python3 deploy/06_generate_failure_data.py\n\n"
    "# Full Omniverse mode (loads USD stage visually):\n"
    "/isaac-sim/python.sh deploy/06_generate_failure_data.py"
)

h3("Failure Scenario Definitions")
add_table(
    ["Failure Type", "Primary Signal", "Secondary Signals", "Real-World Cause", "Onset Speed"],
    [
        ["Overheating",        "temp_c +25 to +45°C",    "power_kw +0.5–1.5 kW\ncpu_load +10–30%",  "CRAC unit failure",      "Rapid (hours)"],
        ["Disk Degradation",   "disk_health -40 to -70%","temp_c +2–8°C\ncpu_load +5%",              "SMART wear, controller", "Slow (weeks)"],
        ["Power Fluctuation",  "power_kw spike +2–5 kW", "temp_c +3–10°C",                           "PDU failure, UPS swap",  "Fast (seconds)"],
        ["Cooling Failure",    "temp_c +15 to +35°C",    "power_kw +1–2.5 kW\ndisk_health -10–30%", "Chiller failure",        "Medium (hours)"],
    ],
    col_widths=[1.5, 1.8, 2.0, 1.8, 1.4]
)

h3("Bell-Curve Ramping")
body(
    "Failures don't appear or disappear instantaneously. The simulation applies a sine "
    "curve to ramp sensor deviations up and then back down, creating natural-looking "
    "anomalies that match real hardware behaviour:"
)
code_block(
    "progress = (step - onset) / (end - onset)  # 0.0 → 1.0\n"
    "ramp     = math.sin(progress * math.pi)    # 0 → 1 → 0  (bell curve)\n"
    "temp    += ramp * random.uniform(temp_delta_low, temp_delta_high)"
)

h3("Dataset Statistics")
add_table(
    ["Metric", "Value", "Notes"],
    [
        ["Total rows",           "414,720",  "48 racks × 8,640 steps (30 days at 5-min intervals)"],
        ["Normal rows",          "401,976",  "96.9% of dataset"],
        ["Total failure rows",   "12,744",   "3.1% across all four types"],
        ["Training samples",     "360,288",  "90% of sliding-window examples"],
        ["Validation samples",   "40,032",   "10% held out"],
        ["CSV file size",        "~25 MB",   "After compression"],
    ],
    col_widths=[2.0, 1.5, 4.0]
)

page_break()

# Phase 7
h2("6.7  Phase 7 — Temporal Transformer World Model")
body(
    "The world model is a Temporal Transformer Encoder with three independent MLP heads, "
    "one per prediction horizon (1 hour, 6 hours, 24 hours ahead). Given the last 60 "
    "minutes of sensor readings for a rack, it outputs failure probabilities for each horizon."
)

h3("Why a Transformer Instead of an LSTM?")
add_table(
    ["Aspect", "LSTM", "Transformer"],
    [
        ["Processing",     "Sequential — step by step",          "Parallel — all steps at once"],
        ["Long-range memory","Degrades with sequence length",    "Full attention over all timesteps"],
        ["Training speed", "Slow (cannot parallelize)",          "Fast (GPU-friendly)"],
        ["Interpretability","Hidden state is opaque",            "Attention weights are inspectable"],
        ["Scale",          "Diminishing returns with depth",     "Scales well with layers + heads"],
    ],
    col_widths=[2.0, 2.5, 3.0]
)

h3("Architecture Details")
code_block(
    "Input:  (batch_size=512, window=12, features=4)\n"
    "        Features: temp_c, power_kw, disk_health, cpu_load\n"
    "        Window  : 12 timesteps × 5 minutes = 60 minutes of history\n"
    "\n"
    "Step 1: Linear projection  4 → d_model=64\n"
    "        (256, 12, 64)\n"
    "\n"
    "Step 2: Positional Encoding  (adds time-order signal)\n"
    "        PE(t, 2i)   = sin(t / 10000^(2i/64))\n"
    "        PE(t, 2i+1) = cos(t / 10000^(2i/64))\n"
    "\n"
    "Step 3: Transformer Encoder  ×3 layers\n"
    "        MultiHeadAttention with 4 heads\n"
    "        d_ff = 128, dropout = 0.1\n"
    "        (256, 12, 64)\n"
    "\n"
    "Step 4: Mean pooling over time dimension\n"
    "        (256, 64)  — single vector per rack\n"
    "\n"
    "Step 5: Three MLP heads (parallel)\n"
    "        Linear(64,32) → ReLU → Dropout → Linear(32,2)\n"
    "        → logits for 1h / 6h / 24h failure probability"
)

h3("Training Configuration")
add_table(
    ["Hyperparameter", "Value", "Reason"],
    [
        ["d_model",        "64",      "Small enough to train fast, large enough to learn patterns"],
        ["Attention heads","4",       "Each head can specialise: temp, power, disk, correlations"],
        ["Encoder layers", "3",       "Sufficient depth without overfitting on ~360K samples"],
        ["Batch size",     "512",     "Fills RTX 3070 VRAM efficiently"],
        ["Learning rate",  "1e-4",    "AdamW with weight decay 1e-4, cosine annealing schedule"],
        ["Epochs",         "15",      "Val loss plateaus around epoch 13–15 on this dataset"],
        ["Window size",    "12 steps","60 minutes of sensor history per prediction"],
    ],
    col_widths=[2.0, 1.2, 4.3]
)

h3("Critical Implementation Details")
body("Two bugs were discovered and fixed during testing that are important for students to understand:")

body("Bug 1 — Missing normalization in predict():", bold=True)
body(
    "The training dataset applies per-window z-score normalization to each sample "
    "(subtract mean, divide by std). The original predict() function passed raw sensor "
    "values directly to the model, which was trained on normalized data. The model saw "
    "completely out-of-distribution inputs and predicted ~1% failure for everything "
    "including a rack at 55°C. Fix: apply the same normalization in predict() before "
    "passing to the model."
)

body("Bug 2 — Single class weight applied to all horizons:", bold=True)
body(
    "Class weights were computed from the 1h horizon label distribution (3.6% positive) "
    "and applied to all three horizons. But the 24h horizon has 17.3% positive rate "
    "(many more windows fall within 24h of a failure), so the 1h weight of 13.8× was "
    "drastically wrong for the 24h head. Fix: compute independent class weights for "
    "each horizon based on its actual positive rate."
)

code_block(
    "# Fixed: per-horizon class weights\n"
    "for horizon_name in HORIZONS:\n"
    "    labels = [dataset.samples[i][1][horizon_name] for i in range(len(dataset))]\n"
    "    n_neg  = sum(1 for l in labels if l == 0)\n"
    "    n_pos  = sum(1 for l in labels if l == 1)\n"
    "    w_neg  = n_total / (2 * n_neg)\n"
    "    w_pos  = n_total / (2 * n_pos)\n"
    "    criteria[horizon_name] = nn.CrossEntropyLoss(\n"
    "        weight=torch.tensor([w_neg, w_pos]))"
)

page_break()

# Phase 8
h2("6.8  Phase 8 — Vertex AI Training Job")
body(
    "Vertex AI is Google's managed machine learning platform. Instead of renting a GPU VM, "
    "installing CUDA, managing storage, and running training manually, you submit a job "
    "specification and Vertex AI handles provisioning, training, and teardown automatically."
)
code_block("source deploy/config.env\npython3 deploy/08_vertex_training.py")
body("The script performs three steps:")
numbered(
    "Packages deploy/07_world_model.py into a Python source distribution and uploads it "
    "to GCS as a training package."
)
numbered(
    "Submits a CustomTrainingJob to Vertex AI specifying: machine type a2-highgpu-1g "
    "(A100 GPU), the PyTorch 2.1 pre-built container, the training data GCS path, and "
    "the output model GCS path. Training typically completes in 20–40 minutes on an A100."
)
numbered(
    "Deploys the trained model to a Vertex AI Endpoint on an n1-standard-4 machine with "
    "a T4 GPU. The endpoint accepts POST requests with sensor windows and returns "
    "failure probabilities as JSON."
)
note(
    "The A100 job costs ~$3/hour. A 40-minute training run costs ~$2. The endpoint "
    "costs ~$1.50/hour while running. Shut down the endpoint when not in use."
)

# Phase 9
h2("6.9  Phase 9 — Inference Integration")
body(
    "The inference configuration (deploy/09_inference_config.toml) connects the Vertex AI "
    "endpoint back to the Omniverse viewer. A polling service reads live sensor data from "
    "GCS every 60 seconds, calls the endpoint for each rack, and writes failure probability "
    "values as USD custom attributes on the corresponding rack prims. The Kit viewer reads "
    "these attributes and renders a red overlay on racks above the alert threshold."
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. STUDENT ASSET SETUP
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Getting the USD Assets (Students)")
body(
    "The 9.6 GB DataHall USD stage is an NVIDIA proprietary asset. Students obtain it "
    "through their instructor via a shared Google Cloud Storage bucket (Option B). "
    "The instructor uploads the asset once and grants each student read-only GCS access."
)

h2("7.1  Instructor Steps (Run Once)")
code_block(
    "# Upload assets to GCS (run once before the class):\n"
    "source deploy/config.env\n"
    "bash deploy/03_upload_assets.sh\n\n"
    "# Grant a single student access:\n"
    "bash deploy/instructor_grant_access.sh student@gmail.com\n\n"
    "# Grant a whole class at once (one email per line in a text file):\n"
    "bash deploy/instructor_grant_access.sh --file class_roster.txt\n\n"
    "# Grant via Google Group (easiest for large classes):\n"
    "bash deploy/instructor_grant_access.sh --group class@googlegroups.com"
)

h2("7.2  Student Steps (Run Once on First Day)")
body("The instructor gives each student: (1) the GitHub repo URL and (2) the GCS bucket name.")
code_block(
    "# Clone the repo:\n"
    "git clone https://github.com/INSTRUCTOR/dc-world-model-tutorial\n"
    "cd dc-world-model-tutorial\n\n"
    "# Run the one-command setup script:\n"
    "GCS_BUCKET=instructor-bucket-name bash deploy/student_setup.sh"
)
body("The student_setup.sh script automatically:")
numbered("Installs gcloud CLI if not present.")
numbered("Opens a browser window for Google account authentication.")
numbered("Verifies the student has access to the bucket (prints a clear error if not).")
numbered("Downloads the full 9.6 GB asset tree to ~/datacenter_assets/.")
numbered("Installs all Python dependencies from requirements.txt.")
numbered("Prints a checklist confirming the student is ready to proceed.")

h2("7.3  Access Control")
add_table(
    ["IAM Role", "Permissions", "Who Gets It"],
    [
        ["storage.objectViewer", "List + download objects. Cannot upload or delete.", "All students"],
        ["storage.objectAdmin",  "Full read/write/delete access.",                   "Instructor only"],
    ],
    col_widths=[2.5, 3.5, 1.5]
)
note("Students cannot modify or delete assets. The instructor's upload is safe.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. TRAINING RESULTS
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Training Results and Model Performance")
body("Hardware used: NVIDIA GeForce RTX 3070 Laptop GPU (8 GB VRAM), CUDA 12.8, PyTorch 2.8.")
body("Training time: approximately 3 minutes for 15 epochs with batch size 512.")

add_table(
    ["Epoch", "Train Loss", "Val Loss", "1h Accuracy", "6h Accuracy", "24h Accuracy"],
    [
        ["1",  "1.9733", "0.6198", "90.0%", "87.8%", "80.3%"],
        ["5",  "1.8197", "0.5915", "86.0%", "84.4%", "75.0%"],
        ["10", "1.8012", "0.5858", "88.4%", "85.9%", "76.6%"],
        ["15", "1.7980", "0.5844", "89.0%", "86.6%", "77.1%"],
    ],
    col_widths=[0.8, 1.3, 1.3, 1.5, 1.5, 1.5]
)

body("Key observations:")
bullet("1h accuracy > 6h > 24h — expected, predicting closer events is easier.")
bullet("Val loss plateaus after ~epoch 13, indicating the model is well-fitted.")
bullet("Class weights applied: 1h=13.8×, 6h=7.5×, 24h=2.9× failure class weighting.")
bullet("Without class weights, all three heads converge to predicting 'normal' for everything.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. LIVE INFERENCE
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Live Inference Results")
body(
    "After training, the model is tested by passing manually constructed sensor windows "
    "representing known scenarios. The predict() function applies z-score normalization "
    "and returns failure probabilities for all three horizons."
)

add_table(
    ["Scenario", "1h Prob.", "6h Prob.", "24h Prob.", "Alert?"],
    [
        ["Normal rack (stable 22°C, healthy disk)",        "31.9%", "41.8%", "47.0%", "No"],
        ["Overheating (temp 28°C → 55°C over 60 min)",    "100.0%","99.9%", "99.8%", "Yes ←"],
        ["Disk degrading (health 0.85 → 0.40)",           "99.9%", "99.9%", "99.6%", "Yes ←"],
        ["Cooling failure (temp 25°C → 48°C + power spike)","100.0%","99.9%","99.9%","Yes ←"],
        ["Power fluctuation (5 kW → 12 kW spike)",        "67.4%", "65.9%", "57.4%", "No"],
    ],
    col_widths=[3.5, 1.0, 1.0, 1.0, 1.0]
)

body("Analysis:")
bullet(
    "Overheating, disk degradation, and cooling failure all trigger 100% alerts because "
    "their sensor signatures are clear, trending patterns that the Transformer can reliably detect."
)
bullet(
    "Power fluctuation reaches 67% — below the 70% alert threshold. This is because the "
    "fluctuation appears in only one feature (power_kw) with no temperature trend, making "
    "it harder to distinguish from normal load spikes."
)
bullet(
    "The normal rack shows 32–47% across horizons. This non-zero baseline reflects the "
    "model's uncertainty when seeing only shape information (absolute values are removed "
    "by per-window normalization)."
)
note(
    "Power fluctuation detection could be improved by: (1) using global normalization "
    "so absolute power values are preserved, or (2) adding a feature for rate-of-change "
    "(delta) of each sensor."
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. EXERCISES
# ══════════════════════════════════════════════════════════════════════════════
h1("10. Exercises for Students")

h2("10.1  Beginner")
numbered("Change FAILURE_RATE from 0.15 → 0.30 in 06_generate_failure_data.py, regenerate data, and retrain. Does accuracy improve? Why or why not?")
numbered("Add a 5th sensor feature fan_rpm to the CSV generator. Modify NUM_FEATURES and retrain. Does the additional signal help?")
numbered("Change WINDOW_SIZE from 12 → 36 (3 hours of history). Measure the change in val accuracy and training time.")
numbered("After training, print a confusion matrix for the 1h horizon: true positives, false positives, true negatives, false negatives.")

h2("10.2  Intermediate")
numbered("Replace mean pooling with a [CLS] token (insert a learnable token at position 0, pool from it instead of averaging). Does accuracy change?")
numbered("Implement early stopping: stop training if val_loss does not improve for 5 consecutive epochs.")
numbered("Write a script that reads inference results and sends a Slack webhook alert when any rack exceeds 80% failure probability.")
numbered("Modify the simulation so cooling failures affect an entire row of racks simultaneously (correlated failures).")

h2("10.3  Advanced")
numbered(
    "Replace z-score normalization with global statistics: compute mean and std over the "
    "entire training set and save them in the model checkpoint. Apply these in predict(). "
    "How does this affect power fluctuation detection?"
)
numbered(
    "Add a Transformer decoder that generates future sensor trajectories (next 12 steps). "
    "This is a generative world model. Use the reconstruction error as an additional "
    "anomaly signal."
)
numbered(
    "Implement federated learning using PySyft or Flower: train one model per simulated "
    "data center without ever centralising raw sensor data. Aggregate gradients only."
)
numbered(
    "Replace the synthetic sensor data with real IPMI readings from a physical server. "
    "How does model performance change when trained on real vs. synthetic data?"
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. GLOSSARY
# ══════════════════════════════════════════════════════════════════════════════
h1("11. Key Concepts Glossary")
add_table(
    ["Term", "Definition"],
    [
        ["USD",                "Universal Scene Description — Pixar/NVIDIA's 3D file format. Scenes are built by composing many referenced sub-files rather than one monolithic file."],
        ["Digital Twin",       "A virtual replica of a physical asset synchronized with real sensor data and capable of running simulations."],
        ["USDC (Crate)",       "Binary compressed USD format. Fast to load, smaller than text USD. The DataHall uses this for geometry files."],
        ["Omniverse Kit",      "NVIDIA's SDK for building real-time 3D applications on USD. Our Kit app renders and streams the DataHall."],
        ["WebRTC",             "Web protocol for real-time video streaming used by Google Meet and Zoom. We use it to stream the 3D viewport to browsers."],
        ["GCS Fuse",           "Mounts a Google Cloud Storage bucket as a Linux filesystem folder. The GPU VM reads USD from GCS as if it were on local disk."],
        ["Transformer Encoder","Neural network architecture using self-attention. Every timestep can attend to every other timestep simultaneously."],
        ["Self-Attention",     "Mechanism that computes a weighted sum of all values, where weights are determined by query-key dot products. Allows learning which timesteps matter."],
        ["Positional Encoding","Sine/cosine signals added to embeddings to give the model a sense of time order (Transformers have no built-in sequence order)."],
        ["World Model",        "A neural network that predicts how a system evolves over time. Our model predicts future failure probability given past sensor readings."],
        ["Class Weights",      "Loss function multipliers inversely proportional to class frequency. Prevents the model from ignoring rare failure events."],
        ["Vertex AI",          "Google's managed ML platform. Handles GPU provisioning, training, and model serving without you managing VMs."],
        ["GCS Bucket",         "A top-level container in Google Cloud Storage. Like an S3 bucket or a hard drive in the cloud."],
        ["Artifact Registry",  "Google's private Docker image registry. Like DockerHub but private and in the same region as your VM for fast pulls."],
        ["Sliding Window",     "Taking a fixed-length contiguous slice of a timeseries as model input — here, the last 12 timesteps (60 minutes)."],
        ["Synthetic Data",     "Data generated by simulation rather than collected from the physical world. Allows creating rare events (failures) at scale."],
    ],
    col_widths=[2.0, 5.5]
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12. COST SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1("12. Cloud Cost Summary")
body("The following table shows estimated costs for running the complete tutorial end to end.")
add_table(
    ["Phase", "Resource", "Duration/Size", "Rate", "Estimated Cost"],
    [
        ["2",   "GCS Bucket (10 GB)",       "1 month",  "$0.02/GB/mo",  "$0.20"],
        ["2",   "Artifact Registry (20 GB)","1 month",  "$0.10/GB/mo",  "$2.00"],
        ["2,5", "GPU VM g2-standard-8 (L4)","8 hours",  "$0.40/hr",     "$3.20"],
        ["3",   "GCS Upload (9.6 GB)",      "One-time", "Free",         "$0.00"],
        ["8",   "Vertex AI A100 Training",  "2 hours",  "$3.00/hr",     "$6.00"],
        ["8",   "Vertex AI Endpoint (T4)",  "4 hours",  "$1.50/hr",     "$6.00"],
        ["-",   "Network egress (~20 GB)",  "One-time", "$0.08/GB",     "$1.60"],
    ],
    col_widths=[0.7, 2.5, 1.8, 1.3, 1.7]
)
body("Total estimated cost per student: approximately $19 for the full tutorial (Phases 1–9).")
body("Cost-saving tips:")
bullet("Stop the GPU VM when not actively streaming: gcloud compute instances stop datacenter-kit-vm --zone us-central1-a")
bullet("Delete the Vertex AI endpoint when not in use — it bills by the hour.")
bullet("Students only need GCS read access, not their own GCP project, to download assets and run Phases 6–7.")
bullet("Set a $25 budget alert in GCP Billing console to avoid surprises.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 13. COMMON ERRORS
# ══════════════════════════════════════════════════════════════════════════════
h1("13. Common Errors and Fixes")
add_table(
    ["Error Message", "Likely Cause", "Fix"],
    [
        ["gcloud: command not found",
         "Phase 1 not complete",
         "Run bash deploy/01_install_gcloud.sh then restart your terminal."],
        ["403 Forbidden / AccessDenied on GCS",
         "Not authenticated or no bucket permission",
         "Run gcloud auth application-default login. Ask instructor to grant access."],
        ["CUDA out of memory",
         "Batch size too large for your GPU",
         "Reduce BATCH_SIZE in 07_world_model.py (try 128 or 256)."],
        ["VM not found",
         "Wrong zone in config.env",
         "Check GCP_ZONE matches where the VM was created."],
        ["docker push denied",
         "Not authenticated to Artifact Registry",
         "Run gcloud auth configure-docker us-central1-docker.pkg.dev"],
        ["Port 8011 connection refused",
         "Kit not started or firewall not open",
         "Check docker logs -f datacenter-kit on the VM. Verify firewall rules."],
        ["[WARN] omni.usd not available",
         "Running with system Python, not Isaac Sim Python",
         "This is normal — the script falls back to standalone CSV-only mode. No action needed."],
        ["KeyError: AIP_TRAINING_DATA_URI",
         "Running 08_vertex_training.py locally",
         "This env var is set by Vertex AI. Submit the job, don't run it locally."],
        ["val_loss not decreasing",
         "Class imbalance — model predicts all normal",
         "Ensure per-horizon class weights are applied in the training loop."],
        ["predict() returns ~1% for all scenarios",
         "Missing normalization in predict()",
         "Apply z-score normalization in predict() before tensor conversion."],
    ],
    col_widths=[2.5, 2.0, 3.0]
)

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
